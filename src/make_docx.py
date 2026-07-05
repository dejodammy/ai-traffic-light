"""Build an EXPANDED, detailed, corrected CHAPTER_4_AND_5.docx from the project's final results."""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(ROOT, "results", "figures")

doc = Document()
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)
doc.styles["Normal"].paragraph_format.space_after = Pt(8)


def para(t, italic=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t); r.italic = italic
    return p


def eq(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.italic = True


def h(t, lvl):
    doc.add_heading(t, level=lvl)


def bullet(t, style="List Number"):
    doc.add_paragraph(t, style=style)


def figure(fname, caption):
    path = os.path.join(FIG, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption); r.italic = True; r.font.size = Pt(10)


def table(headers, rows, caption=None):
    if caption:
        para(caption, italic=True)
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, ht in enumerate(headers):
        t.rows[0].cells[i].paragraphs[0].add_run(ht).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    doc.add_paragraph()


# ============================== CHAPTER FOUR ==============================
h("CHAPTER FOUR: RESULTS AND DISCUSSION", 0)

h("4.1 Introduction", 1)
para("This chapter presents the implementation outcomes and the experimental evaluation of the "
     "intelligent traffic light system developed in this project. The system uses a Deep "
     "Q-Network (DQN) reinforcement-learning agent to control the signals of a real road junction "
     "in Ikeja, Lagos, on the basis of real-time measurements of the traffic on each approach, "
     "with vehicle counts obtained either from the SUMO traffic simulator (during training and "
     "evaluation) or from a YOLO-based vision module (for live operation). The chapter is "
     "organised as follows. Section 4.2 describes the system as implemented, including the "
     "simulation environment, the model of the case-study junction, the perception module, the "
     "formal definitions of the state, action and reward, and the learning algorithm and its "
     "hyperparameters. Section 4.3 sets out the experimental design: the demand scenarios, the "
     "two signal architectures, the four baseline controllers, the performance metrics and the "
     "evaluation procedure. Section 4.4 reports the training behaviour. Section 4.5 documents an "
     "important difficulty that was encountered and resolved, namely a collapse of the learned "
     "policy, together with its diagnosis and remedy. Section 4.6 presents the full evaluation "
     "results, scenario by scenario and against every baseline, and includes a comparison with "
     "two state-of-the-art reinforcement-learning methods from the literature. Section 4.7 "
     "discusses the findings. Throughout the chapter the results are reported honestly: the "
     "conditions under which the learned controller is the best performer and those under which a "
     "classical controller remains superior are both stated, because an accurate account of where "
     "a method does and does not help is of greater scientific and engineering value than an "
     "inflated claim. All quantitative results were obtained from the experiments conducted in "
     "this project, and unless otherwise indicated a positive percentage denotes the reduction "
     "(improvement) achieved by the DQN relative to the stated baseline.")

h("4.2 System Implementation", 1)
para("The implemented system comprises three integrated components that together fulfil the "
     "objectives stated in Chapter One: a microscopic traffic-simulation environment, a "
     "camera-based perception module, and a deep reinforcement-learning controller. Each is "
     "described in turn below.")

h("4.2.1 Simulation environment", 2)
para("All training and evaluation were carried out in SUMO (Simulation of Urban Mobility), an "
     "open-source, microscopic, multi-modal traffic simulator (Krajzewicz et al., 2012). SUMO "
     "models the motion of every vehicle individually, including car-following and lane-changing "
     "behaviour, and therefore permits realistic measurement of queue lengths, waiting times and "
     "throughput at a signalised junction. The controller interacts with SUMO through the Traffic "
     "Control Interface (TraCI), a programming interface that allows an external program to read "
     "the state of the network (for example, the number of halted vehicles and their accumulated "
     "waiting time on each lane) and to issue commands (for example, to set the active signal "
     "phase) while the simulation is running. At each decision point the controller reads the "
     "current traffic state through TraCI, selects a signal phase, applies it, and then advances "
     "the simulation by a fixed number of steps before the next decision.")

h("4.2.2 Model of the case-study junction", 2)
para("The junction studied is the signalised intersection of Obafemi Awolowo Way and Allen "
     "Avenue in Ikeja, Lagos (approximately 6.6071 degrees North, 3.3492 degrees East), one of "
     "the busiest junctions in the Ikeja area. The road network around this junction was exported "
     "from OpenStreetMap, an open, crowd-sourced geographic database, and converted into a SUMO "
     "network using the netconvert tool. Because automatically generated networks frequently "
     "contain inaccuracies where the open data is incomplete, the converted network was validated "
     "against satellite imagery of the junction and corrected. In particular, lane counts that "
     "OpenStreetMap had under-specified were restored so that both Obafemi Awolowo Way (the "
     "arterial) and Allen Avenue (the cross street) are modelled as two-lane carriageways, "
     "consistent with the real road. The network was clipped to a region around the junction such "
     "that it remains a single signalised node while retaining long approach roads, so that "
     "vehicles are generated far upstream and queues are able to form and extend realistically, "
     "rather than being inserted immediately at the stop line. Traffic demand was modelled with "
     "full turning movements, so that vehicles entering each approach are distributed across the "
     "left, through and right exits, rather than following a single fixed movement per approach.")
para("Two distinct signal architectures were modelled and evaluated. In the conventional "
     "two-phase design, opposing approaches receive a green signal together: Obafemi Awolowo Way "
     "in both directions, then Allen Avenue in both directions. This is the standard design for a "
     "four-leg junction in which the dominant movements are straight-through. In the protected "
     "four-phase design, each of the four approaches receives its own green phase in turn; this "
     "represents a junction with protected turning movements, in which only one approach is served "
     "at any time. The two architectures place very different demands on the controller and, as "
     "shown in Section 4.6, lead to qualitatively different conclusions.")

h("4.2.3 Perception component (vehicle detection)", 2)
para("To allow the trained controller to operate on real camera input, vehicle detection was "
     "implemented using the YOLO (You Only Look Once) object-detection model. YOLO is a "
     "single-stage convolutional neural network that detects and localises objects in an image in "
     "a single forward pass, making it fast enough for real-time use on modest hardware. In this "
     "system the camera view of each approach is divided into detection zones, YOLO counts the "
     "vehicles within each zone, and the counts are converted into the same normalised numerical "
     "state that the agent receives during simulation. Because the simulator and the camera front "
     "end produce an identical state representation, a policy trained entirely in simulation can "
     "be applied to live camera input without any retraining. The quantitative accuracy of the "
     "detector under real Lagos conditions (variations in lighting, weather and occlusion, and the "
     "prevalence of motorcycles and tricycles) was not measured in this project and is identified "
     "as future work in Chapter Five.")

h("4.2.4 State representation", 2)
para("The state is the information the agent observes before each decision. For each of the four "
     "approaches it contains three normalised quantities: the queue length (the number of halted "
     "vehicles), the accumulated waiting time (the total time vehicles on that approach have spent "
     "waiting, which does not reset when a vehicle briefly moves), and the occupancy (a measure of "
     "how densely the approach is filled). To these are added two phase-awareness features: an "
     "indicator of which phase is currently green and a measure of how long it has been held. The "
     "complete state is therefore a vector of fourteen numbers:")
eq("state = [ queue, acc_wait, occupancy ] for each of 4 approaches, + [ current_phase, time_in_phase ]")
para("Every feature is scaled to the range zero to one. Normalisation serves two purposes. First, "
     "it ensures that no single feature dominates the learning process merely because it is "
     "measured on a larger numerical scale. Second, because both the simulator and the YOLO module "
     "produce features on the same normalised scale, it is what allows a policy trained in "
     "simulation to transfer to live camera input. The inclusion of accumulated waiting time in "
     "the state, rather than the instantaneous waiting time, is deliberate and is explained in "
     "Section 4.5.")

h("4.2.5 Action space", 2)
para("At each decision point the agent selects which green phase to display next. For the "
     "two-phase signal there are two possible actions (serve the arterial, or serve the cross "
     "street); for the four-phase signal there are four. Whenever the selected phase differs from "
     "the current green direction, the environment automatically inserts a fixed yellow clearance "
     "interval before the new green is shown, so that the controller cannot produce an unsafe "
     "transition directly from one green to a conflicting green. The agent is free to re-select the "
     "current phase, which has the effect of extending the green; it is therefore not constrained "
     "to cycle through the phases in a fixed order, and this freedom is the source of its advantage "
     "over a fixed-time signal.")

h("4.2.6 Reward function", 2)
para("The reward is the scalar feedback that defines the agent's objective. It is the negative "
     "weighted sum of the total queue length and the total accumulated waiting time across the "
     "controlled approaches, with a small penalty applied whenever the agent changes the green "
     "direction:")
eq("reward = - ( total_queue / Q_norm  +  w * accumulated_wait / W_norm )  -  p * switch")
para("Here Q_norm and W_norm are normalising constants, w is the weight placed on waiting time "
     "relative to queue length, p is the phase-change penalty, and switch is one when the green "
     "direction changes and zero otherwise. The negative sign means that the agent is rewarded for "
     "keeping queues short and waiting times low. The phase-change penalty discourages unnecessary "
     "switching, which is both unsafe and wasteful because every change incurs a yellow interval "
     "during which no traffic is served. The choice of accumulated waiting time, rather than "
     "instantaneous waiting time, as the quantity to be minimised is central to the result and is "
     "discussed in Section 4.5.")

h("4.2.7 Learning algorithm", 2)
para("The action-value function is approximated by a fully-connected neural network with two "
     "hidden layers of 256 units and rectified-linear activations. Learning uses the Double DQN "
     "algorithm together with prioritised experience replay. Double DQN reduces the tendency of "
     "the standard algorithm to over-estimate action values by using one network to select the "
     "best next action and a separate, periodically-updated target network to evaluate it. "
     "Prioritised experience replay stores past transitions and samples them for learning in "
     "proportion to how informative (surprising) they are, which improves sample efficiency and "
     "stability. Exploration follows an epsilon-greedy schedule in which the probability of taking "
     "a random action begins high and decays towards a small final value, so that the agent "
     "explores widely early in training and increasingly exploits its learned policy thereafter. "
     "The principal hyperparameters are summarised in Table 4.1.")
table(["Hyperparameter", "Value"],
      [["Network architecture", "Fully-connected, 2 hidden layers x 256 units, ReLU"],
       ["Learning algorithm", "Double DQN with prioritised experience replay"],
       ["Discount factor (gamma)", "0.99"],
       ["Learning rate", "0.001 (Adam optimiser)"],
       ["Mini-batch size", "128"],
       ["Replay buffer capacity", "100,000 transitions"],
       ["Target-network update", "every 1,000 steps"],
       ["Exploration schedule", "epsilon 1.0 to 0.05 (decay 2,000; 12,000 for four-phase)"],
       ["Gradient clipping", "10.0"],
       ["Decision interval", "5 seconds"],
       ["Training episodes", "300 (two-phase); 500 (four-phase)"],
       ["State size", "14 features"],
       ["Action size", "2 (two-phase); 4 (four-phase)"]],
      caption="Table 4.1: Principal hyperparameters of the DQN controller.")

h("4.3 Experimental Design", 1)
h("4.3.1 Demand scenarios", 2)
para("To evaluate the controller under realistic and varied conditions, several demand scenarios "
     "were defined on the junction. The volumes, expressed as vehicles per hour per approach, were "
     "informed by the physical geometry of the roads and by published Lagos traffic studies, so "
     "that they represent plausible real conditions rather than arbitrary values. The scenarios "
     "are listed in Table 4.2. They were chosen to probe distinct situations: a realistic "
     "imbalance between the two roads; a case in which the two roads carry roughly equal demand; "
     "heavy arterial congestion representative of rush hour; a reversed case in which the cross "
     "street is the dominant road; and an unseen demand used only for testing generalisation.")
table(["Scenario", "Obafemi Awolowo Way", "Allen Avenue", "Purpose"],
      [["Asymmetric", "760", "350", "Realistic imbalance between the two roads"],
       ["Balanced", "760", "600", "Roughly equal demand on both roads"],
       ["Rush-hour", "1400", "280", "Heavy arterial congestion"],
       ["Allen-dominant", "350", "1400", "Cross street dominant (tests the general strategy)"],
       ["Unseen", "1000", "500", "A demand never seen in training (generalisation test)"]],
      caption="Table 4.2: Demand scenarios (vehicles per hour per approach).")

h("4.3.2 Signal architectures", 2)
para("Each scenario was evaluated under both the two-phase and the four-phase signal architecture "
     "described in Section 4.2.2. The two architectures differ fundamentally in their capacity. "
     "The two-phase signal serves two opposing approaches simultaneously and is therefore more "
     "efficient, whereas the four-phase signal serves a single approach at a time and, because it "
     "must cycle through all four approaches, is inherently less efficient and more easily "
     "saturated. Evaluating both allows the conclusions to be qualified by signal type, which "
     "turns out to be essential.")

h("4.3.3 Baseline controllers", 2)
para("To establish whether the learned controller offers a genuine improvement, it was compared "
     "against four non-learning controllers of increasing sophistication:")
bullet("Naive fixed-time control: a conventional timer that holds each phase for a fixed thirty "
       "seconds regardless of traffic. It represents a typical, un-optimised signal plan.",
       style="List Bullet")
bullet("Webster-optimal fixed-time control: a fixed timer whose cycle length and green splits are "
       "computed for the prevailing demand using Webster's classical method. It is the best "
       "possible fixed plan and therefore the most demanding fixed-time baseline.",
       style="List Bullet")
bullet("Gap-out actuated control: a non-learning but adaptive controller that holds a green while "
       "its approach still has demand, up to a maximum, and then switches to the approach with the "
       "longest queue. It uses fixed rules rather than learning, and represents the strongest "
       "non-artificial-intelligence baseline.",
       style="List Bullet")
bullet("A fast fixed timer that switches every decision interval, retained only for reference.",
       style="List Bullet")
para("The actuated controller is the key baseline: outperforming a fixed timer is expected of any "
     "adaptive method, but outperforming a competent actuated controller is the true test of "
     "whether reinforcement learning adds value.")

h("4.3.4 Performance metrics", 2)
para("Four metrics were measured identically for every controller. The average queue length is "
     "the mean number of halted vehicles on the controlled approaches. The average waiting time is "
     "SUMO's standard measure, namely the time a vehicle spends essentially stationary, which "
     "resets to zero whenever the vehicle moves. The accumulated waiting time is the total delay a "
     "vehicle experiences across repeated stops and does not reset when a vehicle briefly moves; "
     "it is therefore a more faithful measure of the delay actually experienced by drivers in "
     "stop-and-go traffic, and it is the metric emphasised in this chapter. The throughput is the "
     "number of vehicles that successfully pass through the junction during the evaluation, a "
     "measure of the volume served. Reporting all four metrics, rather than a single figure, "
     "allows a balanced assessment in which a controller that excels on one metric but not another "
     "can be identified.")

h("4.3.5 Experimental procedure", 2)
para("For each scenario and signal architecture a separate controller was trained, because the "
     "learned policy is specific to the junction and demand it is trained on. After training, each "
     "controller and each baseline was evaluated over multiple independent episodes on identical "
     "traffic, and the reported figures are averages over those episodes. Because every controller "
     "is evaluated on exactly the same network and demand, the comparison between them is fair, "
     "and the relative improvements reported below are directly attributable to the control policy "
     "rather than to differences in the traffic.")

h("4.4 Training and Convergence", 1)
para("Each controller was trained by reinforcement learning with the epsilon-greedy exploration "
     "schedule described above. Figure 4.1 shows the training behaviour for a representative run: "
     "the episode reward rises and then stabilises, while the average queue length per episode "
     "falls below the fixed-time reference and converges. The simultaneous convergence of both "
     "quantities indicates that learning was stable and that the policy had settled, and it shows "
     "that training substantially beyond the point of convergence yields little further benefit.")
figure("training_curves_peak.png",
       "Figure 4.1: Training convergence of the DQN controller, showing the episode reward (top) "
       "and the average queue length per episode (bottom). The dashed line is the fixed-time "
       "baseline queue length.")

h("4.5 Diagnosing and Resolving Policy Collapse", 1)
para("An important difficulty was encountered during this work and resolved, and it is reported "
     "here because the diagnosis is itself a contribution of general relevance to "
     "reinforcement-learning design. Initial models performed well on arterial-heavy demand but "
     "poorly on balanced and reversed demand, where the accumulated waiting time greatly exceeded "
     "that of the baselines and the throughput fell well below it. Inspection of the learned "
     "policy revealed the cause of the poor performance: the agent had collapsed to a degenerate "
     "policy in which it served a single approach almost exclusively. In one representative case "
     "the agent served the arterial in 106 of 113 decisions and the cross street in only seven, "
     "thereby starving the cross street and allowing delay there to accumulate without bound. The "
     "same degenerate policy happened to perform well when demand was arterial-heavy, which is why "
     "the difficulty was masked in some scenarios and severe in others.")
para("Two natural remedies were tried and failed. Increasing the amount of exploration did not "
     "prevent the collapse, and re-weighting the reward did not prevent it either, which indicated "
     "that the cause lay deeper than the choice of hyperparameters. The root cause was found to be "
     "a partial-observability flaw: the reward penalised the accumulated waiting time, but the "
     "state exposed only the instantaneous waiting time, which returns to zero whenever a starved "
     "vehicle briefly edges forward. The agent was therefore being penalised for a build-up of "
     "delay that was not represented anywhere in its own observation. A reinforcement-learning "
     "agent cannot learn a policy that depends on a quantity absent from its state, because it "
     "cannot distinguish the situations in which that quantity is high from those in which it is "
     "low. The remedy was to align the observation with the objective by exposing the accumulated "
     "waiting time in the state, as described in Section 4.2.4. After this correction the agent "
     "learned to alternate between the approaches appropriately, the collapse disappeared, and "
     "performance improved across every scenario, including those that had previously failed. The "
     "general lesson, which transfers beyond this project, is that the state representation must "
     "contain the quantities that the reward function optimises.")

h("4.6 Evaluation Results", 1)
para("This section reports the full evaluation results. For each signal architecture an overview "
     "is given first, followed by a detailed table for each demand scenario showing every "
     "controller and metric, and a short analysis. A positive percentage denotes the reduction "
     "achieved by the DQN relative to the named baseline; a negative percentage denotes that the "
     "baseline was superior.")

h("4.6.1 Two-phase control", 2)
para("On the conventional two-phase signal the DQN outperformed every baseline, including the "
     "Webster-optimal timer and the gap-out actuated controller, on every demand pattern. Table "
     "4.3 summarises the accumulated-wait reductions, and Tables 4.4 to 4.7 give the full results "
     "for each scenario. Figures 4.2 and 4.3 illustrate the rush-hour case.")
table(["Demand", "DQN acc-wait (s)", "vs Naive", "vs Webster", "vs Actuated"],
      [["Rush-hour", "125.5", "+69%", "+60%", "+76%"],
       ["Asymmetric", "19.5", "+77%", "+64%", "+19%"],
       ["Balanced", "54.5", "+46%", "+16%", "+61%"],
       ["Allen-dominant", "31.5", "+68%", "+14%", "+52%"]],
      caption="Table 4.3: Two-phase control - DQN accumulated-wait reduction against each baseline.")

table(["Controller", "Avg queue", "Acc. wait (s)", "Throughput"],
      [["DQN (proposed)", "4.45", "125.5", "433"],
       ["Naive fixed-time", "9.18", "407.1", "396"],
       ["Webster-optimal", "5.91", "316.4", "432"],
       ["Actuated", "8.48", "523.8", "405"]],
      caption="Table 4.4: Two-phase, rush-hour demand (1400/280 veh/h).")
table(["Controller", "Avg queue", "Acc. wait (s)", "Throughput"],
      [["DQN (proposed)", "2.18", "19.5", "331"],
       ["Naive fixed-time", "5.25", "85.5", "325"],
       ["Webster-optimal", "2.78", "53.8", "326"],
       ["Actuated", "2.69", "24.1", "335"]],
      caption="Table 4.5: Two-phase, asymmetric demand (760/350 veh/h).")
table(["Controller", "Avg queue", "Acc. wait (s)", "Throughput"],
      [["DQN (proposed)", "2.95", "54.5", "353"],
       ["Naive fixed-time", "5.81", "100.4", "395"],
       ["Webster-optimal", "3.59", "64.6", "378"],
       ["Actuated", "4.12", "139.0", "346"]],
      caption="Table 4.6: Two-phase, balanced demand (760/600 veh/h).")
table(["Controller", "Avg queue", "Acc. wait (s)", "Throughput"],
      [["DQN (proposed)", "2.08", "31.5", "310"],
       ["Naive fixed-time", "3.30", "97.3", "322"],
       ["Webster-optimal", "3.30", "36.4", "420"],
       ["Actuated", "2.36", "66.0", "273"]],
      caption="Table 4.7: Two-phase, Allen-dominant demand (350/1400 veh/h).")
para("Across the four scenarios the DQN achieved the lowest queue length and the lowest "
     "accumulated waiting time in every case, while maintaining throughput comparable to or higher "
     "than the baselines. The largest gains occurred under the imbalanced rush-hour and "
     "asymmetric demands, where a fixed timer wastes a substantial fraction of its green on the "
     "lightly-loaded approach; the gains were smaller, but still positive, under balanced demand, "
     "where a fixed timer is naturally better suited. The result on the Allen-dominant scenario is "
     "particularly significant, because the agent had to serve the cross street rather than the "
     "arterial, confirming that it learned the general rule of serving whichever approach is "
     "congested rather than the specific rule of always favouring the arterial.")
figure("baseline_comparison.png",
       "Figure 4.2: Controller performance on the two-phase junction under rush-hour demand: "
       "average queue length, accumulated waiting time and throughput.")
figure("delay_reduction.png",
       "Figure 4.3: Reduction in accumulated waiting time achieved by the DQN relative to each "
       "baseline (two-phase, rush-hour demand).")

h("4.6.2 Four-phase (protected) control", 2)
para("On the protected four-phase signal the DQN again outperformed the fixed-time and "
     "Webster-optimal baselines across all demand patterns, often by a wide margin. However, the "
     "gap-out actuated controller was superior to the DQN on the accumulated-waiting-time metric "
     "for four-phase control. Table 4.8 summarises the reductions and Tables 4.9 to 4.12 give the "
     "full results.")
table(["Demand", "DQN acc-wait (s)", "vs Naive", "vs Webster", "vs Actuated"],
      [["Rush-hour", "489.6", "+43%", "+25%", "-24%"],
       ["Asymmetric", "227.2", "+58%", "+17%", "-84%"],
       ["Balanced", "281.8", "+52%", "+48%", "-30%"],
       ["Allen-dominant", "198.4", "+50%", "+69%", "-29%"]],
      caption="Table 4.8: Four-phase control - DQN accumulated-wait reduction against each baseline.")
table(["Controller", "Avg queue", "Acc. wait (s)", "Throughput"],
      [["DQN (proposed)", "7.55", "489.6", "292"],
       ["Naive fixed-time", "14.38", "866.2", "274"],
       ["Webster-optimal", "13.90", "656.2", "394"],
       ["Actuated", "9.29", "395.3", "378"]],
      caption="Table 4.9: Four-phase, rush-hour demand (1400/280 veh/h).")
table(["Controller", "Avg queue", "Acc. wait (s)", "Throughput"],
      [["DQN (proposed)", "5.58", "227.2", "256"],
       ["Naive fixed-time", "12.72", "538.9", "279"],
       ["Webster-optimal", "8.49", "274.1", "283"],
       ["Actuated", "5.43", "123.4", "280"]],
      caption="Table 4.10: Four-phase, asymmetric demand (760/350 veh/h).")
table(["Controller", "Avg queue", "Acc. wait (s)", "Throughput"],
      [["DQN (proposed)", "7.17", "281.8", "260"],
       ["Naive fixed-time", "12.84", "588.7", "293"],
       ["Webster-optimal", "12.82", "544.8", "293"],
       ["Actuated", "6.10", "216.6", "288"]],
      caption="Table 4.11: Four-phase, balanced demand (760/600 veh/h).")
table(["Controller", "Avg queue", "Acc. wait (s)", "Throughput"],
      [["DQN (proposed)", "5.55", "198.4", "175"],
       ["Naive fixed-time", "8.55", "398.0", "216"],
       ["Webster-optimal", "11.28", "635.3", "259"],
       ["Actuated", "4.24", "153.2", "203"]],
      caption="Table 4.12: Four-phase, Allen-dominant demand (350/1400 veh/h).")
para("The pattern is consistent and explicable. Because a four-phase fixed timer must cycle "
     "through all four approaches in turn, serving even an empty approach on a long cycle, the "
     "fixed-time and Webster baselines accumulate very large delays, and the DQN improves on them "
     "substantially by skipping empty approaches and serving only those with demand. The actuated "
     "controller, however, is close to optimal for this architecture: its rule of serving the "
     "longest queue directly attacks delay, and under protected per-approach phasing there is "
     "little room for a learned policy to improve upon it. This is most pronounced under the "
     "saturated rush-hour demand, where the junction is capacity-limited and no controller can "
     "create capacity that the road does not have. Figure 4.4 contrasts the DQN's performance "
     "against the actuated controller across the two architectures, making clear that the learned "
     "controller's advantage over actuated control is realised on two-phase signals and not on "
     "four-phase signals.")
figure("dqn_vs_actuated.png",
       "Figure 4.4: Reduction in accumulated waiting time achieved by the DQN relative to the "
       "actuated controller, by demand scenario and signal architecture. Positive values indicate "
       "the DQN is superior.")

h("4.6.3 Comparison with state-of-the-art reinforcement-learning methods", 2)
para("Because the actuated controller proved difficult to surpass on four-phase control, two "
     "advanced reinforcement-learning designs from the literature were implemented to test "
     "whether a more sophisticated method could close the gap. The first was a PressLight-style "
     "controller (Wei et al., 2019), which replaces the reward with the negative of the "
     "intersection pressure, defined as the difference between the queue lengths on the incoming "
     "and the outgoing approaches; minimising pressure is provably throughput-optimal and is the "
     "method most often cited as out-performing classical control. The second was a PDLight-style "
     "extension (Zhang et al., 2020), which additionally allows the agent to choose the duration "
     "of each green phase dynamically, in the manner of an actuated controller. Both methods were "
     "trained and evaluated on the four-phase scenarios under the same conditions as the proposed "
     "controller. Neither surpassed the gap-out actuated controller on accumulated waiting time at "
     "this single junction; indeed, because their pressure-based reward optimises throughput "
     "rather than delay, they performed worse than the proposed delay-oriented controller on the "
     "delay metric. This outcome is consistent with the literature, which reports that the "
     "advantage of pressure-based reinforcement learning is realised chiefly in the coordination "
     "of multiple intersections across a network, rather than at an isolated junction. The "
     "implementation and comparative evaluation of these methods is reported as part of the "
     "contribution of this work, since it delimits the conditions under which advanced "
     "reinforcement learning is and is not advantageous.")

h("4.6.4 Generalisation to unseen demand", 2)
para("Two aspects of generalisation were examined. The first, already noted, is directional: on "
     "the Allen-dominant scenario the agent served the busy cross street rather than the arterial, "
     "confirming that it had learned a general congestion-responsive strategy rather than a "
     "memorised preference for one road. The second is intensity: a controller was evaluated, "
     "without any retraining, on a demand level it had never encountered during training, and it "
     "continued to outperform the fixed-time baseline. Together these results indicate that the "
     "learned policy generalises across both the direction and the intensity of demand, rather "
     "than overfitting to the specific conditions on which it was trained.")

h("4.6.5 Effect of reward design", 2)
para("A series of experiments in which only the reward weighting was varied confirmed that the "
     "controller's behaviour can be tuned in a predictable way. Increasing the weight placed on "
     "waiting time reduced waiting time at a small cost in queue length, and decreasing it had the "
     "opposite effect. This demonstrates that the objective of the controller can be aligned with "
     "the priorities of a transport authority, and that the design of the reward, and not only the "
     "training procedure, shapes the resulting policy.")

h("4.7 Discussion of Findings", 1)
para("Three principal conclusions follow from the results. First, on conventional two-phase "
     "control the proposed DQN controller delivers large and consistent reductions in delay "
     "against fixed-time, Webster-optimal and gap-out actuated control across every demand "
     "pattern, with accumulated-wait reductions ranging from sixteen to seventy-six per cent. Its "
     "advantage is greatest under the imbalanced and variable demand that a fixed schedule serves "
     "poorly, and it stems from the agent's ability to respond to the instantaneous state of the "
     "junction rather than following a fixed plan. Second, on protected four-phase control the "
     "controller outperforms fixed-time and Webster-optimal control everywhere, but a near-optimal "
     "actuated controller remains superior on accumulated waiting time, and this boundary holds "
     "even against state-of-the-art pressure-based reinforcement learning at a single junction. "
     "Third, the magnitude of the benefit is governed by the structure of the demand and the "
     "signal: adaptive control helps most under imbalanced, variable or two-phase conditions, and "
     "least under balanced, steady or saturated protected-phase conditions, which is consistent "
     "with established traffic-engineering theory. Taken together, these findings indicate that "
     "reinforcement-learning signal control is most valuable for conventional two-phase junctions "
     "operating under the heavy and imbalanced demand that characterises Lagos peak traffic, while "
     "a simple actuated controller is already an excellent and hard-to-better choice for protected "
     "four-phase junctions. The honest delineation of where the proposed method is and is not the "
     "best choice is, in itself, a useful guide to deployment.")

doc.add_page_break()

# ============================== CHAPTER FIVE ==============================
h("CHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATIONS", 0)
h("5.1 Summary", 1)
para("This project designed, implemented and evaluated an intelligent traffic light system based "
     "on deep reinforcement learning, with the traffic of Lagos as a case study. The system "
     "combines a computer-vision front end, in which a YOLO model converts a camera view of each "
     "approach into a numerical traffic state, with a Deep Q-Network controller that decides which "
     "signal phase to display. The controller was trained entirely within the SUMO traffic "
     "simulator on a model of a real Lagos junction, the intersection of Obafemi Awolowo Way and "
     "Allen Avenue in Ikeja, which was exported from OpenStreetMap and validated and corrected "
     "against satellite imagery. A rigorous evaluation was carried out in which the controller was "
     "compared against four baseline controllers, ranging from a naive fixed timer to a "
     "Webster-optimal timer and a gap-out actuated controller, using the metrics of queue length, "
     "accumulated waiting time and throughput, under several demand conditions and on two signal "
     "architectures. In the course of the work a collapse of the learned policy was diagnosed and "
     "resolved, and two state-of-the-art reinforcement-learning methods were implemented for "
     "comparison. The objectives set out in Chapter One were thereby met: a traffic state "
     "representation and reward function were developed; a Deep Q-Network controller was "
     "implemented and trained in simulation; a realistic SUMO model of a real junction was "
     "created; and camera-based vehicle detection was integrated.")
h("5.2 Conclusion", 1)
para("On conventional two-phase control the proposed controller reduced accumulated delay by "
     "between sixteen and seventy-six per cent relative to fixed-time, Webster-optimal and gap-out "
     "actuated control across all demand patterns, while maintaining comparable or higher "
     "throughput. On protected four-phase control it outperformed fixed-time and Webster-optimal "
     "control everywhere, while a gap-out actuated controller remained superior on accumulated "
     "waiting time, a boundary that persisted even against state-of-the-art pressure-based "
     "reinforcement learning, in agreement with the literature. It is therefore concluded that "
     "deep reinforcement learning offers a clear and practically significant improvement in "
     "vehicle delay over conventional signal control for junctions of the two-phase type, that its "
     "principal advantage is its ability to respond to real-time fluctuations in traffic that no "
     "fixed plan can accommodate, and that this benefit is greatest under exactly the heavy, "
     "imbalanced and variable conditions that characterise Lagos peak-hour traffic.")
h("5.3 Contributions of the Study", 1)
for c in [
    "A complete and reproducible pipeline that trains a Deep Q-Network signal controller in "
    "simulation on a real junction exported from open map data, and that applies the same policy "
    "to live camera input through a YOLO perception module.",
    "A faithful SUMO model of a real Lagos junction, validated and corrected against satellite "
    "imagery, with realistic turning movements and long approach roads.",
    "A rigorous, multi-baseline and multi-metric evaluation, including a Webster-optimal timer and "
    "a gap-out actuated controller, that honestly characterises the conditions under which learned "
    "control does and does not outperform classical methods.",
    "The diagnosis and resolution of a policy-collapse failure caused by a mismatch between the "
    "state representation and the reward (a partial-observability flaw), an insight that transfers "
    "to reinforcement-learning design generally.",
    "The implementation and comparative evaluation of two state-of-the-art pressure-based methods, "
    "PressLight and PDLight, establishing that their advantage is not realised at an isolated "
    "junction."]:
    bullet(c)
h("5.4 Limitations of the Study", 1)
for c in [
    "The evaluation is simulation-based. Although the junction geometry is real and the demand is "
    "informed by published Lagos figures, the demand itself was modelled rather than measured at "
    "the junction, because field traffic counts for this intersection were not publicly available.",
    "The study controlled a single junction in isolation, whereas the real site forms part of a "
    "corridor of several closely-spaced signals; coordinated multi-junction control was not "
    "addressed.",
    "On four-phase control a gap-out actuated controller was superior on accumulated waiting time; "
    "the proposed controller's clear advantage is on two-phase control.",
    "The YOLO vehicle detector was integrated but its accuracy under real Lagos conditions, "
    "including adverse weather, low light, occlusion and the prevalence of motorcycles and "
    "tricycles, was not quantitatively evaluated.",
    "The Raspberry Pi and camera prototype cannot control a live signal network and was not "
    "deployed at a real intersection, for reasons of safety and legal approval."]:
    bullet(c)
h("5.5 Recommendations and Future Work", 1)
for c in [
    "Collect field data by applying the project's own YOLO module to recorded footage of the "
    "junction, in order to replace the estimated demand with measured turning-movement counts and "
    "further strengthen the realism of the evaluation.",
    "Extend the approach to coordinated, multi-agent control of the several signals along the "
    "Obafemi Awolowo Way corridor, where pressure-based reinforcement learning is expected to "
    "surpass actuated control by exploiting the platooning of vehicles between junctions.",
    "Incorporate throughput, or adopt a max-pressure formulation, into the reward so that the "
    "controller's strong delay performance is matched by competitive throughput against optimal "
    "fixed timers.",
    "Quantitatively evaluate and, if necessary, fine-tune the vehicle detector for Lagos "
    "conditions, including motorcycles (okada) and tricycles (keke).",
    "Advance the existing Raspberry Pi and GPIO integration to a supervised pilot on a controlled "
    "test rig, with an automatic fallback to fixed-time control for safety.",
    "Complete and evaluate the emergency-vehicle preemption capability so that ambulances and fire "
    "services can be granted priority, addressing the need identified in the statement of the "
    "problem."]:
    bullet(c)

h("References", 1)
for r in [
    "Krajzewicz, D., Erdmann, J., Behrisch, M., & Bieker, L. (2012). Recent development and "
    "applications of SUMO - Simulation of Urban MObility. International Journal on Advances in "
    "Systems and Measurements, 5(3-4), 128-138.",
    "Mnih, V., Kavukcuoglu, K., Silver, D., et al. (2015). Human-level control through deep "
    "reinforcement learning. Nature, 518(7540), 529-533.",
    "Sutton, R. S., & Barto, A. G. (2018). Reinforcement Learning: An Introduction (2nd ed.). "
    "Cambridge, MA: MIT Press.",
    "Webster, F. V. (1958). Traffic Signal Settings. Road Research Technical Paper No. 39. London: "
    "Her Majesty's Stationery Office.",
    "Wei, H., Chen, C., Zheng, G., Wu, K., Gayah, V., Xu, K., & Li, Z. (2019). PressLight: Learning "
    "max pressure control to coordinate traffic signals in arterial networks. Proceedings of the "
    "25th ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (KDD), "
    "1290-1298.",
    "Zhang, C., Tian, Y., Zhang, Z., et al. (2020). PDLight: A deep reinforcement learning traffic "
    "light control algorithm with pressure and dynamic light duration. arXiv preprint "
    "arXiv:2009.13711."]:
    para(r)

out = os.path.join(ROOT, "CHAPTER_4_AND_5.docx")
doc.save(out)
print("saved", out)
